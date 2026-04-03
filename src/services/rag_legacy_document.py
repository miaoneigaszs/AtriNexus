"""RAG legacy 文档解析与切块辅助函数。"""

from __future__ import annotations

import logging
import os
import re
from typing import Any, Dict, List

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger("wecom")


def parse_document(file_path: str) -> str:
    """根据后缀名智能提取文档文本。"""
    if not os.path.exists(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return _parse_pdf(file_path)
        if ext in [".docx", ".doc"]:
            return _parse_docx(file_path)
        if ext in [".txt", ".md", ".csv", ".json"]:
            with open(file_path, "r", encoding="utf-8") as file_obj:
                return file_obj.read()

        logger.warning(f"暂未支持的文件格式解析: {ext}")
        return ""
    except Exception as exc:
        logger.error(f"解析文档 {file_path} 失败: {exc}")
        return ""


def _parse_pdf(file_path: str) -> str:
    """解析 PDF 文档，优先使用 TOC，无 TOC 时按字体大小推断标题。"""
    if not fitz:
        logger.error("未安装 PyMuPDF，无法解析 PDF")
        return ""

    try:
        doc = fitz.open(file_path)
        toc = doc.get_toc()

        if not toc:
            final_text = _extract_pdf_headers_by_font(doc)
            total_pages = len(doc)
            doc.close()
            logger.debug(f"PDF解析完成(字体识别): {file_path}, 总页数={total_pages}")
            return final_text

        bookmark_pages: Dict[int, List[tuple[int, str]]] = {}
        for level, title, page_num, *_ in toc:
            bookmark_pages.setdefault(page_num, []).append((level, title))

        result_lines: List[str] = []
        total_pages = len(doc)
        for page_idx, page in enumerate(doc):
            page_num = page_idx + 1
            for level, title in bookmark_pages.get(page_num, []):
                md_level = min(max(level, 1), 6)
                result_lines.append(f"{'#' * md_level} {title}")
                result_lines.append("")

            page_text = page.get_text().strip()
            if page_text:
                result_lines.append(page_text)
            result_lines.append("")

        doc.close()
        logger.debug(f"PDF解析完成: {file_path}, 书签数={len(toc)}, 总页数={total_pages}")
        return "\n".join(result_lines)
    except Exception as exc:
        logger.error(f"PDF解析异常: {file_path}, 错误: {exc}")
        return ""


def _extract_pdf_headers_by_font(doc) -> str:
    """当 PDF 没有书签时，尝试通过字体大小识别标题。"""
    try:
        result_lines: List[str] = []
        for page in doc:
            text_dict = page.get_text("dict")
            if not text_dict or "blocks" not in text_dict:
                continue

            text_blocks: List[Dict[str, Any]] = []
            for block in text_dict["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        text_blocks.append(
                            {
                                "text": text,
                                "size": span.get("size", 12),
                                "bold": "bold" in span.get("font", "").lower(),
                                "y": span.get("bbox", [0, 0, 0, 0])[1],
                            }
                        )

            if not text_blocks:
                continue

            avg_size = sum(block["size"] for block in text_blocks) / len(text_blocks)
            text_blocks.sort(key=lambda item: item["y"])

            for block in text_blocks:
                text = block["text"]
                size = block["size"]
                if size > avg_size * 1.3 or (block["bold"] and size > avg_size * 1.1):
                    if size > avg_size * 2:
                        result_lines.append(f"# {text}")
                    elif size > avg_size * 1.6:
                        result_lines.append(f"## {text}")
                    elif size > avg_size * 1.3:
                        result_lines.append(f"### {text}")
                    else:
                        result_lines.append(f"#### {text}")
                else:
                    result_lines.append(text)
            result_lines.append("")
        return "\n".join(result_lines)
    except Exception as exc:
        logger.warning(f"PDF字体标题识别失败: {exc}")
        return ""


def _parse_docx(file_path: str) -> str:
    """解析 Word 文档并尽量保留标题层级。"""
    if not Document:
        logger.error("未安装 python-docx，无法解析 Word")
        return ""

    try:
        doc = Document(file_path)
        result_lines: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading ") or style_name.startswith("标题 "):
                try:
                    level = min(max(int(style_name.split()[-1]), 1), 6)
                    result_lines.append(f"{'#' * level} {text}")
                except (ValueError, IndexError):
                    result_lines.append(text)
            elif style_name.lower() in ["title", "标题"]:
                result_lines.append(f"# {text}")
            else:
                result_lines.append(text)

        logger.debug(f"DOCX解析完成: {file_path}, 段落数={len(doc.paragraphs)}")
        return "\n".join(result_lines)
    except Exception as exc:
        logger.error(f"DOCX解析异常: {file_path}, 错误: {exc}")
        return ""


def recursive_character_split(text: str, chunk_size: int = 600, overlap: int = 100) -> List[str]:
    """递归分层切块。"""
    separators = ["\n\n", "\n", "。", "！", "？", ".", " "]
    return _do_split(text, separators, chunk_size, overlap)


def _do_split(text: str, separators: List[str], chunk_size: int, overlap: int) -> List[str]:
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    separator = separators[0]
    for sep in separators:
        if sep == "":
            separator = sep
            break
        if sep in text:
            separator = sep
            break

    splits = text.split(separator) if separator else list(text)

    chunks: List[str] = []
    current_chunk: List[str] = []
    current_length = 0

    for piece in splits:
        piece_len = len(piece) + (len(separator) if separator else 0)
        if current_length + piece_len > chunk_size and current_chunk:
            chunk_str = separator.join(current_chunk)
            if chunk_str.strip():
                chunks.append(chunk_str)

            overlap_chars = 0
            keep_pieces: List[str] = []
            for current_piece in reversed(current_chunk):
                if overlap_chars + len(current_piece) <= overlap:
                    keep_pieces.insert(0, current_piece)
                    overlap_chars += len(current_piece) + (len(separator) if separator else 0)
                else:
                    break
            current_chunk = keep_pieces
            current_length = overlap_chars

        current_chunk.append(piece)
        current_length += piece_len

    if current_chunk:
        chunk_str = separator.join(current_chunk)
        if chunk_str.strip():
            chunks.append(chunk_str)

    final_chunks: List[str] = []
    next_seps = separators[1:] if len(separators) > 1 else [""]
    for chunk in chunks:
        if len(chunk) > chunk_size and next_seps:
            final_chunks.extend(_do_split(chunk, next_seps, chunk_size, overlap))
        else:
            final_chunks.append(chunk)
    return final_chunks


def split_markdown(text: str) -> List[Dict[str, Any]]:
    """将 Markdown 文本按标题层级拆分，带上 Header 元数据。"""
    lines = text.split("\n")
    chunks: List[Dict[str, Any]] = []
    current_headers: Dict[str, str] = {}
    current_content: List[str] = []
    header_regex = re.compile(r"^(#{1,6})\s+(.*)")

    for line in lines:
        match = header_regex.match(line)
        if match:
            if current_content:
                content_str = "\n".join(current_content).strip()
                if content_str:
                    chunks.append({"content": content_str, "headers": dict(current_headers)})
                current_content = []

            level = len(match.group(1))
            header_text = match.group(2).strip()
            keys_to_remove = [key for key in current_headers if int(key.replace("H", "")) >= level]
            for key in keys_to_remove:
                del current_headers[key]
            current_headers[f"H{level}"] = header_text
            continue

        current_content.append(line)

    if current_content:
        content_str = "\n".join(current_content).strip()
        if content_str:
            chunks.append({"content": content_str, "headers": dict(current_headers)})

    return chunks

"""
知识库检索增强引擎 (RAG Engine)
负责处理上传的 PDF、Word、TXT、MD 文件。
实现：文本提取、高级递归切块 (Chunking)、向量化入库、元数据过滤搜索。
"""

import os
import re
import json
import logging
import hashlib
import asyncio
from typing import List, Dict, Any, Optional
from datetime import datetime
from collections import defaultdict

import chromadb

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    from docx import Document
except ImportError:
    Document = None

import jieba
from rank_bm25 import BM25Okapi

from src.services.ai.embedding_service import EmbeddingService
from src.utils.async_utils import run_sync
from data.config import config

logger = logging.getLogger('wecom')


class RAGEngine:
    """知识库检索增强引擎"""

    def __init__(self, chroma_client=None):
        # 初始化持久化 ChromaDB Client
        if chroma_client:
            self.client = chroma_client
        else:
            db_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
                'data', 'vectordb'
            )
            os.makedirs(db_path, exist_ok=True)
            self.client = chromadb.PersistentClient(path=db_path)
            
        # BM25 Sparse 检索需要的语料库及分词后索引（内存缓寸，启动时自动从 Chroma 重建）
        # 结构: user_id -> {"corpus": [list of strings], "bm25": BM25Okapi object, "ids": [list of ids]}
        self._bm25_store = {}

        # 初始化共享 Embedding 服务（使用独立的 embedding 配置）
        self._embedding_service = EmbeddingService()
        # 优先使用 embedding_settings，回退到 llm_settings
        if hasattr(config, 'embedding') and config.embedding and config.embedding.api_key:
            embed_key = config.embedding.api_key
            embed_url = config.embedding.base_url or 'https://api.siliconflow.cn/v1'
            logger.info(f"使用 embedding_settings 配置: base_url={embed_url}")
        else:
            embed_key = config.llm.api_key
            embed_url = config.llm.base_url
            logger.warning("embedding_settings 未配置，回退使用 llm_settings（可能不支持 embedding API）")
        if embed_key:
            self._embedding_service.configure(api_key=embed_key, base_url=embed_url)

        if self._embedding_service.is_available():
            self.embedding_fn = self._embedding_service.embedding_function
            self.reranker = self._embedding_service.reranker
        else:
            logger.warning("未配置 API Key，RAG 引擎不可用")
            self.embedding_fn = None
            self.reranker = None

        # 可配置参数（未来可从 config.json 读取）
        self.batch_size = getattr(config.behavior, 'rag_batch_size', 8)  # 批量写入大小，避免 413 错误
        self.chunk_size = getattr(config.behavior, 'rag_chunk_size', 800)  # 文档切块大小
        self.chunk_overlap = getattr(config.behavior, 'rag_chunk_overlap', 150)  # 切块重叠字符数

    def _calculate_batch_size(self, chunks: List[str]) -> int:
        """
        智能计算批量大小，基于文档块平均长度动态调整
        
        API 限制约 8K tokens，按平均字符数估算：
        - 大块 (>2000字符): 4 个/批
        - 中块 (1000-2000字符): 6 个/批
        - 小块 (<1000字符): 8 个/批
        """
        if not chunks:
            return self.batch_size
        
        total_chars = sum(len(c) for c in chunks)
        avg_chunk_size = total_chars / len(chunks)
        
        if avg_chunk_size > 2000:
            return 4
        elif avg_chunk_size > 1000:
            return 6
        return self.batch_size  # 默认 8

    def _get_kb_collection(self, user_id: str):
        """获取专属知识库 Collection，按用户隔离"""
        if not self.embedding_fn:
            return None
        collection_name = f"kb_{user_id}".replace("-", "_")
        try:
            return self.client.get_or_create_collection(
                name=collection_name,
                embedding_function=self.embedding_fn,
                metadata={"description": "User Knowledge Base"}
            )
        except Exception as e:
            logger.error(f"获取知识库集合失败: {e}")
            return None

    def _init_or_update_bm25(self, user_id: str, collection=None):
        """初始化或更新用户的 BM25 索引"""
        if collection is None:
            collection = self._get_kb_collection(user_id)
        if not collection:
            return
            
        try:
            # 取出所有文档
            all_data = collection.get(include=["documents"])
            docs = all_data.get("documents", [])
            ids = all_data.get("ids", [])
            
            if not docs:
                self._bm25_store[user_id] = None
                return
                
            tokenized_corpus = [list(jieba.cut(doc)) for doc in docs]
            bm25 = BM25Okapi(tokenized_corpus)
            self._bm25_store[user_id] = {
                "corpus": docs,
                "ids": ids,
                "bm25": bm25
            }
            logger.debug(f"BM25 索引初始化/更新完成: 用户 {user_id}, {len(docs)} 个 Chunk")
        except Exception as e:
            logger.error(f"初始化 BM25 失败: {e}")

    # ---------- 文档解析阶段 ----------

    def parse_document(self, file_path: str) -> str:
        """根据后缀名智能提取文档文本"""
        if not os.path.exists(file_path):
            return ""

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif ext in ['.txt', '.md', '.csv', '.json']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                logger.warning(f"暂未支持的文件格式解析: {ext}")
                return ""
        except Exception as e:
            logger.error(f"解析文档 {file_path} 失败: {e}")
            return ""

    def _parse_pdf(self, file_path: str) -> str:
        """
        解析 PDF 文档，支持目录结构识别

        策略：
        1. 获取 PDF 书签目录 (TOC)
        2. 使用 get_text("dict") 获取带坐标的文本块
        3. 将书签转换为 Markdown 标题，插入到对应位置
        4. 如果没有书签，尝试通过字体大小/样式识别标题
        """
        if not fitz:
            logger.error("未安装 PyMuPDF，无法解析 PDF")
            return ""

        try:
            doc = fitz.open(file_path)
            toc = doc.get_toc()  # [[level, title, page_num], ...]

            # 如果没有书签，尝试通过字体大小识别标题
            if not toc:
                final_text = self._extract_pdf_headers_by_font(doc, file_path)
                doc.close()
                logger.debug(f"PDF解析完成(字体识别): {file_path}, 总页数={len(doc)}")
                return final_text

            # 构建 书签 -> 页面内容的映射
            bookmark_pages = {}  # page_num -> [(level, title)]
            for item in toc:
                level, title, page_num = item[0], item[1], item[2]
                if page_num not in bookmark_pages:
                    bookmark_pages[page_num] = []
                bookmark_pages[page_num].append((level, title))

            result_lines = []
            total_pages = len(doc)

            for page_idx, page in enumerate(doc):
                page_num = page_idx + 1  # PDF 页码从1开始

                # 如果该页有书签，在页面开头插入 Markdown 标题
                if page_num in bookmark_pages:
                    for level, title in bookmark_pages[page_num]:
                        # 转换为 Markdown 标题（限制层级 1-6）
                        md_level = min(max(level, 1), 6)
                        result_lines.append(f"{'#' * md_level} {title}")
                        result_lines.append("")  # 空行分隔

                # 获取页面文本
                page_text = page.get_text().strip()
                if page_text:
                    result_lines.append(page_text)
                result_lines.append("")  # 页面间空行分隔

            doc.close()
            final_text = "\n".join(result_lines)

            logger.debug(f"PDF解析完成: {file_path}, 书签数={len(toc)}, 总页数={total_pages}")
            return final_text

        except Exception as e:
            logger.error(f"PDF解析异常: {file_path}, 错误: {e}")
            return ""

    def _extract_pdf_headers_by_font(self, doc, file_path: str) -> str:
        """
        当 PDF 没有书签时，尝试通过字体大小识别标题
        使用 get_text("dict") 获取带坐标和字体信息的文本块
        """
        try:
            result_lines = []
            
            for page in doc:
                # 获取带格式的文本块
                text_dict = page.get_text("dict")
                if not text_dict or "blocks" not in text_dict:
                    continue
                
                # 收集所有文本块及其字体大小
                text_blocks = []
                for block in text_dict["blocks"]:
                    if block.get("type") != 0:  # 跳过图片块
                        continue
                    
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue
                            font_size = span.get("size", 12)
                            is_bold = "bold" in span.get("font", "").lower()
                            text_blocks.append({
                                "text": text,
                                "size": font_size,
                                "bold": is_bold,
                                "y": span.get("bbox", [0, 0, 0, 0])[1]  # Y坐标
                            })
                
                if not text_blocks:
                    continue
                
                # 计算平均字体大小，识别标题（字体较大或加粗）
                avg_size = sum(b["size"] for b in text_blocks) / len(text_blocks)
                
                # 按 Y 坐标排序，保持阅读顺序
                text_blocks.sort(key=lambda x: x["y"])
                
                for block in text_blocks:
                    text = block["text"]
                    size = block["size"]
                    
                    # 判断是否为标题：字体显著大于平均值 或 加粗且较大
                    if size > avg_size * 1.3 or (block["bold"] and size > avg_size * 1.1):
                        # 根据字体大小确定标题层级
                        if size > avg_size * 2:
                            result_lines.append(f"# {text}")
                        elif size > avg_size * 1.6:
                            result_lines.append(f"## {text}")
                        elif size > avg_size * 1.3:
                            result_lines.append(f"### {text}")
                        else:
                            result_lines.append(f"#### {text}")
                    else:
                        result_lines.append(text)
                
                result_lines.append("")  # 页面分隔
            
            return "\n".join(result_lines)
            
        except Exception as e:
            logger.warning(f"PDF字体标题识别失败: {e}")
            return ""

    def _parse_docx(self, file_path: str) -> str:
        """
        解析 Word 文档，支持标题样式识别
        
        策略：
        1. 检测段落样式名（Heading 1-6）
        2. 转换为 Markdown 标题格式
        3. 复用 split_markdown 进行结构化切块
        """
        if not Document:
            logger.error("未安装 python-docx，无法解析 Word")
            return ""
        
        try:
            doc = Document(file_path)
            result_lines = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 检测标题样式
                style_name = para.style.name if para.style else ""
                
                # 匹配 Heading 1-6 或 标题 1-6
                if style_name.startswith("Heading ") or style_name.startswith("标题 "):
                    try:
                        # 提取标题级别
                        level_str = style_name.split()[-1]
                        level = int(level_str)
                        level = min(max(level, 1), 6)  # 限制 1-6
                        result_lines.append(f"{'#' * level} {text}")
                    except (ValueError, IndexError):
                        result_lines.append(text)
                elif style_name.lower() in ["title", "标题"]:
                    # 文档主标题
                    result_lines.append(f"# {text}")
                else:
                    # 普通段落
                    result_lines.append(text)
            
            final_text = "\n".join(result_lines)
            logger.debug(f"DOCX解析完成: {file_path}, 段落数={len(doc.paragraphs)}")
            return final_text
            
        except Exception as e:
            logger.error(f"DOCX解析异常: {file_path}, 错误: {e}")
            return ""

    # ---------- 智能切块阶段 (Chunking) ----------

    def recursive_character_split(self, text: str, chunk_size: int = 600, overlap: int = 100) -> List[str]:
        """
        递归分层切块法：优先按照段落(\n\n)切分子句，再按换行(\n)，最后按句号切分。
        保留完整的句子上下文结构。
        """
        separators = ["\n\n", "\n", "。", "！", "？", ".", " "]
        return self._do_split(text, separators, chunk_size, overlap)

    def _do_split(self, text: str, separators: List[str], chunk_size: int, overlap: int) -> List[str]:
        # base case
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        separator = separators[0]
        # 寻找当前适用的最佳分隔符
        for sep in separators:
            if sep == "":
                separator = sep
                break
            if sep in text:
                separator = sep
                break

        # 尝试拆分
        if separator:
            splits = text.split(separator)
        else:
            splits = list(text)  # 退化成单字符

        chunks = []
        current_chunk = []
        current_length = 0

        for s in splits:
            s_len = len(s) + (len(separator) if separator else 0)
            if current_length + s_len > chunk_size and current_chunk:
                # 合并现有的部分作为一个块
                chunk_str = separator.join(current_chunk)
                if chunk_str.strip():
                    chunks.append(chunk_str)

                # 处理重叠
                overlap_chars = 0
                keep_pieces = []
                for p in reversed(current_chunk):
                    if overlap_chars + len(p) <= overlap:
                        keep_pieces.insert(0, p)
                        overlap_chars += len(p) + (len(separator) if separator else 0)
                    else:
                        break
                current_chunk = keep_pieces
                current_length = overlap_chars

            current_chunk.append(s)
            current_length += s_len

        if current_chunk:
            chunk_str = separator.join(current_chunk)
            if chunk_str.strip():
                chunks.append(chunk_str)

        # 如果切出来的块依然太大，递归使用下一级分隔符
        final_chunks = []
        next_seps = separators[1:] if len(separators) > 1 else [""]
        for c in chunks:
            if len(c) > chunk_size and next_seps:
                final_chunks.extend(self._do_split(c, next_seps, chunk_size, overlap))
            else:
                final_chunks.append(c)

        return final_chunks

    def split_markdown(self, text: str) -> List[Dict[str, Any]]:
        """
        将 Markdown 文本按标题层级拆分，带上 Header 元数据。
        非 Markdown 文本视为一个整体大块。
        """
        lines = text.split('\n')
        chunks = []
        current_headers = {}
        current_content = []

        header_regex = re.compile(r'^(#{1,6})\s+(.*)')

        for line in lines:
            match = header_regex.match(line)
            if match:
                # 遇到新标题，先保存上一个块的内容
                if current_content:
                    content_str = '\n'.join(current_content).strip()
                    if content_str:
                        chunks.append({"content": content_str, "headers": dict(current_headers)})
                    current_content = []

                level = len(match.group(1))
                header_text = match.group(2).strip()

                # 清理同级或更低级别的旧标题
                keys_to_remove = [k for k in current_headers.keys() if int(k.replace('H', '')) >= level]
                for k in keys_to_remove:
                    del current_headers[k]

                current_headers[f'H{level}'] = header_text
            else:
                current_content.append(line)

        # 最后一个块
        if current_content:
            content_str = '\n'.join(current_content).strip()
            if content_str:
                chunks.append({"content": content_str, "headers": dict(current_headers)})

        return chunks

    # ---------- 核心接口：入库与检索 ----------

    def add_document(self, user_id: str, file_name: str, file_path: str, category: str = "默认分类") -> tuple[bool, str]:
        """
        完整的一键入库闭环流：读取文件 -> 解析格式 -> MD5防重校验 -> 结构化拆分 -> 递归切块 -> 入库
        """
        collection = self._get_kb_collection(user_id)
        if not collection:
            return False, "向量库初始化失败"

        # 1. 提取文本
        text = self.parse_document(file_path)
        if not text.strip():
            return False, "解析出的文本为空，暂不支持纯图片或损坏的文件"

        # 2. 文档级 MD5 防污防重校验
        doc_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
        try:
            exist_results = collection.get(where={"doc_hash": doc_hash}, limit=1)
            if exist_results and exist_results['ids']:
                logger.info(f"文档防重触发: {file_name} (hash={doc_hash}) 已存在，跳过入库。")
                try:
                    os.remove(file_path)
                except:
                    pass
                return True, "该文档内容已存在于您的知识库中，无需重复投喂！"
        except Exception as e:
            logger.warning(f"防重校验异常，继续入库: {e}")

        # 3. 结构化拆分 -> 基础层级感知
        structured_sections = self.split_markdown(text)

        # 4. 对每个结构块进行带 Overlap 的智能切片
        final_chunks = []
        for section in structured_sections:
            sub_chunks = self.recursive_character_split(section["content"], chunk_size=self.chunk_size, overlap=self.chunk_overlap)
            for sub in sub_chunks:
                final_chunks.append({
                    "content": sub,
                    "headers": section["headers"]
                })

        if not final_chunks:
            return False, "文件文本分割失败"

        # 5. 准备写入 ChromaDB
        ids = []
        documents = []
        metadatas = []
        upload_time = datetime.now().isoformat()

        for i, item in enumerate(final_chunks):
            chunk_content = item["content"]
            headers_meta = item["headers"]

            ids.append(f"kb_{doc_hash}_chunk_{i}")
            documents.append(chunk_content)

            # 将 Header 原文也拼接到摘要中，保证语义丰富
            meta_dict = {
                "file_name": file_name,
                "category": category,
                "chunk_index": i,
                "upload_time": upload_time,
                "doc_hash": doc_hash
            }
            # 把 H1, H2 直接铺平放入 metadata 供后续精准检索
            meta_dict.update(headers_meta)
            metadatas.append(meta_dict)

        # 6. 执行写入（智能分批，避免 413 错误）
        try:
            total_chunks = len(ids)
            smart_batch_size = self._calculate_batch_size(documents)
            
            for i in range(0, total_chunks, smart_batch_size):
                batch_ids = ids[i:i + smart_batch_size]
                batch_docs = documents[i:i + smart_batch_size]
                batch_metas = metadatas[i:i + smart_batch_size]
                
                collection.add(
                    ids=batch_ids,
                    documents=batch_docs,
                    metadatas=batch_metas
                )
                logger.debug(f"批次 {i//smart_batch_size + 1}/{(total_chunks-1)//smart_batch_size + 1} 写入完成")
            
            logger.info(f"文档入库成功: {file_name}, Hash={doc_hash}, 共切分 {total_chunks} 块, 批大小={smart_batch_size}")
            # 更新 BM25 索引
            self._init_or_update_bm25(user_id, collection)
            try:
                os.remove(file_path)
            except:
                pass
            return True, f"成功存入知识库，为 AI 构建了 {total_chunks} 个记忆神经元。"
        except Exception as e:
            logger.error(f"知识块写入数据库异常: {e}")
            return False, f"存入失败: {str(e)}"

    def retrieve_knowledge(self, user_id: str, query: str, top_k: int = 3,
                           category_filter: Optional[str] = None,
                           h1_filter: Optional[str] = None,
                           h2_filter: Optional[str] = None,
                           skip_rerank: bool = False) -> List[Dict]:
        """
        检索知识库（支持分类和标题层级过滤）

        流程：
        1. 向量相似度粗筛（召回 top_k * 5，默认查 15 条）
        2. Reranker 重排序（Cross-Encoder 精打分，提取最终 Top-K）

        Args:
            user_id: 用户ID
            query: 查询文本
            top_k: 返回结果数量
            category_filter: 分类过滤
            h1_filter: H1 标题过滤
            h2_filter: H2 标题过滤
            skip_rerank: 跳过 Reranker，直接返回向量相似度结果（用于快速判断）
        """
        collection = self._get_kb_collection(user_id)
        if not collection:
            return []

        if collection.count() == 0:
            return []

        # 构建复合过滤条件
        where_clause = {}
        if category_filter:
            where_clause["category"] = category_filter
        if h1_filter:
            where_clause["H1"] = h1_filter
        if h2_filter:
            where_clause["H2"] = h2_filter

        # 确定检索数量：快速模式只取 top_k，海选模式多取一些
        fetch_k = top_k if skip_rerank else min(max(top_k * 5, 15), collection.count())
        
        query_args = {
            "query_texts": [query],
            "n_results": fetch_k
        }
        if where_clause:
            query_args["where"] = where_clause

        try:
            fetch_k_dense = fetch_k
            fetch_k_sparse = fetch_k

            # === 并行双路召回 (Dense + Sparse) ===
            import concurrent.futures

            dense_results_raw = {}
            sparse_results_raw = {}
            
            # 使用 ThreadPoolExecutor 并行发起 Dense(Chroma) 和 Sparse(BM25) 查询
            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                def run_dense():
                    return collection.query(**query_args)

                def run_sparse():
                    if user_id not in self._bm25_store:
                        # 如未初始化，先初始化
                        self._init_or_update_bm25(user_id, collection)
                    
                    bm25_data = self._bm25_store.get(user_id)
                    if not bm25_data:
                        return []
                        
                    tokenized_query = list(jieba.cut(query))
                    scores = bm25_data["bm25"].get_scores(tokenized_query)
                    
                    # 按照分数降序排列，取前 fetch_k_sparse 名
                    top_n_idx = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:fetch_k_sparse]
                    
                    sparse_res = []
                    for idx in top_n_idx:
                        if scores[idx] > 0: # 过滤掉得分为0的
                            sparse_res.append((bm25_data["ids"][idx], scores[idx]))
                    return sparse_res

                future_dense = executor.submit(run_dense)
                future_sparse = executor.submit(run_sparse)
                
                # 获取结果
                try:
                    dense_results_raw = future_dense.result(timeout=10)
                except Exception as e:
                    logger.error(f"Dense 向量查询失败: {e}")
                    
                try:
                    sparse_results_raw = future_sparse.result(timeout=10)
                except Exception as e:
                    logger.error(f"Sparse BM25查询失败: {e}")

            # === 解析与融合 (RRF) ===
            rrf_k = 60 # RRF 算法的平滑常数
            rrf_scores = defaultdict(float)
            
            # 记录 id -> object (由于我们需要返回带 metadata 的结果，所以要把查到的 doc 和 meta 全部映射起来)
            doc_meta_map = {}

            # 处理 Dense 榜单，注入到 RRF
            if dense_results_raw and dense_results_raw.get('documents') and dense_results_raw['documents'][0]:
                docs = dense_results_raw['documents'][0]
                ids = dense_results_raw['ids'][0]
                metas = dense_results_raw['metadatas'][0] if dense_results_raw.get('metadatas') else [{}] * len(docs)
                distances = dense_results_raw['distances'][0] if dense_results_raw.get('distances') else [0] * len(docs)
                
                for rank, (doc_id, doc, meta, dist) in enumerate(zip(ids, docs, metas, distances)):
                    # Dense RRF Score
                    rrf_scores[doc_id] += 1.0 / (rrf_k + rank + 1)
                    
                    # 记录实体
                    if doc_id not in doc_meta_map:
                        doc_meta_map[doc_id] = {"content": doc, "metadata": meta, "dense_dist": dist}

            # 处理 Sparse 榜单，注入到 RRF
            if sparse_results_raw:
                # Sparse 是 [(id, score), ...]
                # 为了拿到 docs 和 metadata，如果 doc_id 不在 map 里，我们需要去 Chroma 反查一次
                missing_ids = [doc_id for doc_id, _ in sparse_results_raw if doc_id not in doc_meta_map]
                if missing_ids:
                    missing_data = collection.get(ids=missing_ids, include=["documents", "metadatas"])
                    if missing_data and missing_data.get("ids"):
                        m_ids = missing_data["ids"]
                        m_docs = missing_data["documents"]
                        m_metas = missing_data["metadatas"]
                        for i in range(len(m_ids)):
                            doc_meta_map[m_ids[i]] = {"content": m_docs[i], "metadata": m_metas[i]}

                for rank, (doc_id, bm25_score) in enumerate(sparse_results_raw):
                    if doc_id in doc_meta_map:
                        # Sparse RRF Score
                        rrf_scores[doc_id] += 1.0 / (rrf_k + rank + 1)
            
            # 依据 RRF 得分排序，截取前 fetch_k 个用于后续重排
            sorted_rrf_ids = sorted(rrf_scores.keys(), key=lambda k: rrf_scores[k], reverse=True)[:fetch_k]
            
            # 组装为最终混合海选列表，同时进行[层级上下文增强]
            fused_docs = []
            fused_metas = []
            fused_original_distances = []
            
            for doc_id in sorted_rrf_ids:
                item = doc_meta_map[doc_id]
                original_text = item["content"]
                meta = item["metadata"]
                dist = item.get("dense_dist", 0.5) # BM25找回的默认给一个居中的欧式距离
                
                # --- 层级上下文增强 ---
                h1 = meta.get("H1", "")
                h2 = meta.get("H2", "")
                h3 = meta.get("H3", "")
                
                # 拼接目录树前缀，让大模型在打分(Rerank)和推理时感知结构
                path_parts = [p for p in [h1, h2, h3] if p]
                if path_parts:
                    path_str = " > ".join(path_parts)
                    enriched_text = f"[所属章节: {path_str}]\n{original_text}"
                else:
                    enriched_text = original_text
                    
                fused_docs.append(enriched_text)
                fused_metas.append(meta)
                fused_original_distances.append(dist)

            if not fused_docs:
                return []

            if skip_rerank:
                logger.debug(f"[快速混合检索] 查询='{query[:20]}...' 返回 {len(fused_docs)} 条，跳过重排")
            else:
                logger.info(f"RAG海选(BM25+Dense融合): '{query[:20]}...' 获取候选 {len(fused_docs)} 块。开始调用大模型二次重排。")

                # 2. 第二阶段：重排模型精确打分 (Reranking)
                if self.reranker and len(fused_docs) > 1:
                    rerank_results = self.reranker.rerank(query=query, texts=fused_docs, top_n=top_k)
                    if rerank_results:
                        final_ret = []
                        for rank_info in rerank_results:
                            idx = rank_info["index"]
                            score = rank_info["relevance_score"]
                            final_ret.append({
                                "content": fused_docs[idx],
                                "metadata": fused_metas[idx],
                                "score": score
                            })
                        logger.info(f"Reranker 重排完成，提取最终 {len(final_ret)} 个高分片段。")
                        return final_ret
                    else:
                        logger.warning("重排接口未返回数据，降级为默认RRF排序")

            # 降级或直接返回：使用融合分数排列
            ret = []
            for i in range(min(top_k, len(fused_docs))):
                ret.append({
                    "content": fused_docs[i],
                    "metadata": fused_metas[i],
                    "score": rrf_scores[sorted_rrf_ids[i]] * 100 # 将 RRF 放大便于展示
                })
            return ret

        except Exception as e:
            logger.error(f"检索知识库失败: {e}")
            return []

    # ---------- 新增：文档结构大纲功能 ----------

    def get_document_outline(self, user_id: str, file_name: Optional[str] = None) -> Dict[str, Any]:
        """
        获取文档的标题结构大纲

        Args:
            user_id: 用户ID
            file_name: 指定文件名，为 None 则返回所有文档的大纲

        Returns:
            Dict: {
                "documents": {
                    "file1.md": {
                        "H1": ["第一章 概述", "第二章 详细说明"],
                        "H2": ["1.1 背景", "1.2 目的"],
                        "category": "产品文档"
                    }
                },
                "categories": ["产品文档", "规章制度"]
            }
        """
        collection = self._get_kb_collection(user_id)
        if not collection or collection.count() == 0:
            return {"documents": {}, "categories": []}

        try:
            # 构建过滤条件
            where_clause = {}
            if file_name:
                where_clause["file_name"] = file_name

            # 获取元数据
            if where_clause:
                res = collection.get(where=where_clause, include=["metadatas"])
            else:
                res = collection.get(include=["metadatas"])

            metas = res.get("metadatas", [])

            # 按文档聚合标题结构
            docs_structure = defaultdict(lambda: {"H1": set(), "H2": set(), "H3": set(), "category": ""})
            categories = set()

            for meta in metas:
                fname = meta.get("file_name", "未知文件")
                categories.add(meta.get("category", "默认分类"))
                docs_structure[fname]["category"] = meta.get("category", "默认分类")

                # 收集各级标题
                for level in [1, 2, 3]:
                    key = f"H{level}"
                    if meta.get(key):
                        docs_structure[fname][key].add(meta[key])

            # 转换为列表并排序
            result = {"documents": {}, "categories": sorted(list(categories))}
            for fname, structure in docs_structure.items():
                result["documents"][fname] = {
                    "category": structure["category"],
                    "H1": sorted(list(structure["H1"])),
                    "H2": sorted(list(structure["H2"])),
                    "H3": sorted(list(structure["H3"]))
                }

            return result

        except Exception as e:
            logger.error(f"获取文档大纲失败: {e}")
            return {"documents": {}, "categories": []}

    def get_knowledge_list(self, user_id: str) -> Dict[str, List[str]]:
        """获取用户当前知识库里拥有的文档一览，按 category 分组"""
        collection = self._get_kb_collection(user_id)
        if not collection or collection.count() == 0:
            return {}

        try:
            res = collection.get(include=["metadatas"])
            metas = res.get("metadatas", [])

            kb_map = {}
            for m in metas:
                cat = m.get("category", "默认分类")
                fname = m.get("file_name", "未知文件")
                if cat not in kb_map:
                    kb_map[cat] = set()
                kb_map[cat].add(fname)

            return {k: list(v) for k, v in kb_map.items()}
        except Exception as e:
            logger.error(f"获取知识库列表失败: {e}")
            return {}

    def delete_document(self, user_id: str, file_name: str) -> bool:
        """根据文件名从知识库移除相关碎片"""
        collection = self._get_kb_collection(user_id)
        if not collection:
            return False

        try:
            collection.delete(where={"file_name": file_name})
            # 文档被删除后同步刷新 BM25 索引
            self._init_or_update_bm25(user_id, collection)
            return True
        except Exception as e:
            logger.error(f"删除文档碎片失败: {e}")
            return False

    # ---------- 新增：检索结果格式化 ----------

    def format_search_results(self, results: List[Dict], include_score: bool = True) -> str:
        """
        将检索结果格式化为易读的分类展示

        Args:
            results: 检索结果列表
            include_score: 是否包含相关度分数

        Returns:
            str: 格式化后的文本
        """
        if not results:
            return "未找到相关内容。"

        # 按文档和标题分组
        grouped = defaultdict(lambda: defaultdict(list))
        for item in results:
            meta = item.get("metadata", {})
            file_name = meta.get("file_name", "未知文件")
            h1 = meta.get("H1", "")
            h2 = meta.get("H2", "")

            header_path = " > ".join(filter(None, [h1, h2])) or "正文"
            grouped[file_name][header_path].append(item)

        # 格式化输出
        output_lines = ["📚 找到以下内容：\n"]

        for file_name, headers in grouped.items():
            output_lines.append(f"📄 《{file_name}》")
            for header_path, items in headers.items():
                if header_path != "正文":
                    output_lines.append(f"   📌 {header_path}")
                for item in items:
                    content = item.get("content", "")[:150].replace("\n", " ")
                    score = item.get("score", 0)
                    if include_score:
                        output_lines.append(f"      • {content}... [相关度: {score:.2f}]")
                    else:
                        output_lines.append(f"      • {content}...")
            output_lines.append("")

        return "\n".join(output_lines)

    # ========== 异步方法（性能优化）==========

    async def retrieve_knowledge_async(
        self, 
        user_id: str, 
        query: str, 
        top_k: int = 3,
        category_filter: Optional[str] = None,
        h1_filter: Optional[str] = None,
        h2_filter: Optional[str] = None
    ) -> List[Dict]:
        """
        异步检索知识库（不阻塞事件循环）
        
        使用线程池执行 ChromaDB 同步操作，适用于高并发场景。
        """
        return await run_sync(
            self.retrieve_knowledge,
            user_id, query, top_k, category_filter, h1_filter, h2_filter
        )

    async def get_knowledge_list_async(self, user_id: str) -> Dict[str, List[str]]:
        """异步获取知识库列表"""
        return await run_sync(self.get_knowledge_list, user_id)

    async def get_document_outline_async(
        self, 
        user_id: str, 
        file_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """异步获取文档大纲"""
        return await run_sync(self.get_document_outline, user_id, file_name)

    async def delete_document_async(self, user_id: str, file_name: str) -> bool:
        """异步删除文档"""
        return await run_sync(self.delete_document, user_id, file_name)
